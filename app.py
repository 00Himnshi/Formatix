from flask import Flask, render_template, request, send_file
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.use(cors())
@app.route("/")
def home():
    return render_template("index.html")

@app.route("/format", methods=["POST"])
def format_file():
    file = request.files["file"]
    if not file or not file.filename.endswith(".txt"):
        return "Invalid file type. Please upload a .txt file.", 400

    input_path = os.path.join(UPLOAD_FOLDER, file.filename)
    output_path = os.path.join(OUTPUT_FOLDER, "formatted_paper.docx")
    file.save(input_path)

    try:
        format_ieee(input_path, output_path)
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        return f"Error processing file: {e}", 500

def format_ieee(input_path, output_path):
    document = Document()
    
    # Helper functions
    def add_title_page():
        title = document.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.add_run("Title of the Paper")
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.name = "Times New Roman"
        title_run.text = title_run.text.title()  
        document.add_paragraph()

    def margins():
        section = document.sections[0]
        section.top_margin = Inches(0.75)  
        section.bottom_margin = Inches(1)  
        section.left_margin = Inches(0.625)  
        section.right_margin = Inches(0.625)

    def level1(text):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.all_caps = True 

    def level2(text):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.italic = True

    def level3(text):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)  
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.font.italic = True 
        if not text.endswith(":"):
            run.text += ":"

    def references(text):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = Pt(-18)

    def abstract(text):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run("Abstract—")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    def author():
        author_name = "Author1"
        author_affiliation = "Affiliation1"
        author_email = "email1@gmail.com"
        author_paragraph = document.add_paragraph()
        author_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        author_run = author_paragraph.add_run(f"{author_name}, {author_affiliation}, {author_email}")
        author_run.font.name = "Times New Roman"
        author_run.font.size = Pt(10)
        document.add_paragraph()

    def keyword(keywords):
        keyword_paragraph = document.add_paragraph()
        keyword_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        keyword_run = keyword_paragraph.add_run("Keywords—")
        keyword_run.italic = True
        keyword_run.font.name = "Times New Roman"
        keyword_run.font.size = Pt(10)
        keyword_run = keyword_paragraph.add_run(f" {keywords}")
        keyword_run.italic = True
        keyword_run.font.name = "Times New Roman"
        keyword_run.font.size = Pt(10)
        keyword_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def main_content(t):
        paragraph = document.add_paragraph(t)
        run = paragraph.runs[0]
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Read and process the input text
    def read_input(file_path):
        add_title_page()
        author()
        margins()

        with open(file_path, "r", encoding="utf-8") as file:
            lines = file.readlines()

        i = -1
        while i < len(lines):
            i += 1
            line = lines[i].strip()
            words = line.split()
            first = words[0].strip()
            
            if first.lower() == "abstract":
                content = ""
                i += 1
                while i < len(lines) and lines[i].strip() != "":
                    line = lines[i]
                    content += line
                    i += 1
                abstract(content)

            elif first.lower() == "keywords":
                content = ""
                i += 1
                while i < len(lines) and lines[i].strip() != "":
                    line = lines[i]
                    content += line
                    i += 1
                keyword(content)

            elif first.lower() == "references":
                content = ""
                i += 1
                while i < len(lines) and lines[i].strip() != "":
                    line = lines[i]
                    content += line
                    i += 1
                references(content)

            else:
                header_level1 = line
                content = ""
                i += 1
                level1(header_level1)

                while i < len(lines) and lines[i].strip() != "":
                    newline = lines[i].strip()
                    word = newline.split()
                    spl = word[0].strip()
                    
                    if "." in spl:
                        header_level2 = newline
                        level2(header_level2)

                    elif ")" in spl:
                        header_level3 = newline
                        level3(header_level3)

                    else:
                        content = lines[i]
                        main_content(content)

                    i += 1
        document.save(output_path)

    read_input(input_path)

if __name__ == "__main__":
    app.run()
